import re
import time
import threading
from docx import Document
from collections import deque
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextLineHorizontal, LTTextBox, LTTextLine, LTFigure
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter


# 生成id
def get_next_id():
    try:
        key_util = KeyUtil()
        return key_util.next_id()
    except Exception as e:
        print(f'Exception occurred: {e}')


class KeyUtil:
    twepoch = 1288834974657
    worker_id_bits = 3
    datacenter_id_bits = 3
    sequence_bits = 8
    max_worker_id = -1 ^ (-1 << worker_id_bits)
    max_datacenter_id = -1 ^ (-1 << datacenter_id_bits)
    worker_id_shift = sequence_bits
    datacenter_id_shift = sequence_bits + worker_id_bits
    timestamp_left_shift = sequence_bits + worker_id_bits + datacenter_id_bits
    sequence_mask = -1 ^ (-1 << sequence_bits)

    def __init__(self, worker_id=0, datacenter_id=0):
        if worker_id > self.max_worker_id or worker_id < 0:
            raise ValueError(f"worker Id can't be greater than {self.max_worker_id} or less than 0")
        if datacenter_id > self.max_datacenter_id or datacenter_id < 0:
            raise ValueError(f"datacenter Id can't be greater than {self.max_datacenter_id} or less than 0")
        self.worker_id = worker_id
        self.datacenter_id = datacenter_id
        self.sequence = 0
        self.last_timestamp = -1
        self.db_key_id_deque = deque()
        self.lock = threading.Lock()
        self.init_db_key_id()

    def next_id(self):
        with self.lock:
            if not self.db_key_id_deque:
                self.init_db_key_id()
            return self.db_key_id_deque.pop()

    def next_id_do(self):
        with self.lock:
            timestamp = self.time_gen()

            if timestamp < self.last_timestamp:
                raise RuntimeError(
                    f"Clock moved backwards. Refusing to generate id for {self.last_timestamp - timestamp} milliseconds")

            if self.last_timestamp == timestamp:
                self.sequence = (self.sequence + 1) & self.sequence_mask
                if self.sequence == 0:
                    timestamp = self.til_next_millis(self.last_timestamp)
            else:
                self.sequence = 0

            self.last_timestamp = timestamp

            return ((timestamp - self.twepoch) << self.timestamp_left_shift) | (
                    self.datacenter_id << self.datacenter_id_shift) | (
                           self.worker_id << self.worker_id_shift) | self.sequence

    @staticmethod
    def til_next_millis(last_timestamp):
        timestamp = KeyUtil.time_gen()
        while timestamp <= last_timestamp:
            timestamp = KeyUtil.time_gen()
        return timestamp

    @staticmethod
    def time_gen():
        return int(time.time() * 1000)

    def init_db_key_id(self):
        id_set = set()
        for _ in range(1000):
            id_set.add(self.next_id_do())
        self.db_key_id_deque.extend(id_set)


# DOCX格式文本段落结构还原
def get_docx_data(path):
    try:
        docx = Document(path)
        content = []
        for paragraph in docx.paragraphs:
            content.append(paragraph.text)
        return content
    except Exception as e:
        return ""


# 结构化PDF文档数据
def get_text_by_miner(file_path):
    with open(file_path, "rb") as fp:
        parser = PDFParser(fp)
        doc = PDFDocument(parser)
        if not doc.is_extractable:
            print(f"{file_path} cannot be extracted")
            return None

        parser.set_document(doc)
        rsrcmgr = PDFResourceManager()
        laparams = LAParams()
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        texts_pages = {}
        index = 0
        for page in PDFPage.create_pages(doc):
            tmp_page = {}
            interpreter.process_page(page)
            # 使用聚合器来获得内容
            layout = device.get_result()
            pdf_shape = (layout.height, layout.width)
            texts_page = []
            for x in layout:
                if isinstance(x, LTTextLineHorizontal) or isinstance(x, LTTextBox) or isinstance(x, LTFigure):
                    n1 = len(x)
                    if isinstance(x, LTTextLineHorizontal) or isinstance(x, LTTextBox):
                        t1 = x.get_text()
                        r1 = (t1, *x.bbox)
                        texts_page.append(r1)
                    else:
                        for m in x:
                            if hasattr(m, 'get_text'):  # 为了确保不出错先判断对象是否具有 get_text()方法
                                t2 = m.get_text()
                                r2 = (t2, *m.bbox)
                                texts_page.append(r2)
                else:
                    pass
            tmp_page['shape'] = pdf_shape
            tmp_page['text'] = texts_page
            texts_pages[index] = tmp_page
            index += 1
        return texts_pages


# PDF文本段落结构还原
def get_pdf_data(file_path):
    article = ''
    texts_pages = get_text_by_miner(file_path)
    for page_index in range(len(texts_pages)):
        text = texts_pages[page_index]['text']
        for line in text:
            line_text = line[0]
            if (re.findall("^[0-9]{1,3}.*?", line_text) and len(line_text) < 6) or len(line_text) < 3:  # 去除页码标识以及空行
                continue
            else:
                if line_text.endswith(' \n'):
                    article += line_text.replace(' \n', '\n')
                else:
                    article += line_text.rstrip('\n')
    article_list = [i for i in article.split('\n') if i != '']
    return article_list

if __name__ == "__main__":
    file = 'text.pdf'
    data = get_text_by_miner(file)
    print(data)
